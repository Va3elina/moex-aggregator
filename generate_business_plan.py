"""
Генерация бизнес-плана / обращения в МосБиржу за льготным тарифом.
Платформа «Фрейм» — аналитика для частных инвесторов.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import copy

TODAY = date.today().strftime("%d.%m.%Y")
YEAR  = date.today().year

# ─────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="AAAAAA"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)

def para_format(para, size=11, bold=False, color=None,
                align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6):
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    for run in para.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))

def add_heading(doc, text, level=1, color="1F3864"):
    para = doc.add_heading(text, level=level)
    para.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    para.paragraph_format.space_after  = Pt(4)
    for run in para.runs:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
        run.font.size = Pt(13 if level == 1 else 11)
    return para

def add_para(doc, text, size=10.5, bold=False, indent=False,
             space_before=0, space_after=5, color=None):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    if indent:
        para.paragraph_format.left_indent = Cm(0.5)
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return para

def add_bullet(doc, text, size=10.5):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(3)
    run = para.add_run(text)
    run.font.size = Pt(size)
    return para

def add_spacer(doc, pt=6):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    para.paragraph_format.line_spacing = Pt(pt)

# ─────────────────────────────────────────────────────────────
# ПОСТРОЕНИЕ ДОКУМЕНТА
# ─────────────────────────────────────────────────────────────

def build(doc: Document):
    # Поля страницы
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(1.5)

    # ══════════════════════════════════════════════════════════
    # ШАПКА
    # ══════════════════════════════════════════════════════════

    to_block = doc.add_paragraph()
    to_block.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    to_block.paragraph_format.space_after = Pt(4)
    r = to_block.add_run("Генеральному директору\nПАО Московская Биржа\n_______________________")
    r.font.size = Pt(10.5)

    add_spacer(doc, 8)

    from_block = doc.add_paragraph()
    from_block.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    from_block.paragraph_format.space_after = Pt(4)
    r = from_block.add_run("от Индивидуального предпринимателя\n[ФИО]\nИНН: [ИНН]\nАдрес: [адрес]\nТел.: [телефон]\nE-mail: [email]")
    r.font.size = Pt(10.5)

    add_spacer(doc, 14)

    # Заголовок
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("ОБРАЩЕНИЕ")
    r.font.size = Pt(14)
    r.font.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(2)
    r = sub.add_run(
        "о предоставлении льготных условий доступа к данным\n"
        "для аналитической платформы «Фрейм»"
    )
    r.font.size = Pt(12)
    r.font.bold = True

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_after = Pt(16)
    r = date_p.add_run(f"г. Москва, {TODAY}")
    r.font.size = Pt(10.5)
    r.font.italic = True

    # ══════════════════════════════════════════════════════════
    # 1. СУТЬ ОБРАЩЕНИЯ
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "1. Суть обращения")

    add_para(doc,
        "Настоящим обращаемся с просьбой о предоставлении льготных условий "
        "доступа к биржевым данным ПАО Московская Биржа для аналитической "
        "платформы «Фрейм», находящейся на стадии разработки и предназначенной "
        "для частных инвесторов — физических лиц.")

    add_para(doc,
        "Конкретные запросы:", bold=True, space_after=3)
    add_bullet(doc,
        "Бесплатный доступ к данным MOEX BI по срочному рынку, рынку акций "
        "и индексам на период 24 (двадцать четыре) месяца с даты заключения договора.")
    add_bullet(doc,
        "Рассрочка оплаты архивных данных (исторические данные по срочному рынку, "
        "рынку акций и индексам) сроком на 36 (тридцать шесть) месяцев равными долями.")

    # ══════════════════════════════════════════════════════════
    # 2. О ПРОЕКТЕ
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "2. О проекте «Фрейм»")

    add_para(doc,
        "«Фрейм» — аналитическая веб-платформа для частных инвесторов, "
        "работающих на российском фондовом рынке. Платформа предоставляет "
        "профессиональные аналитические инструменты, ранее доступные "
        "преимущественно институциональным участникам, в формате удобном "
        "для розничного инвестора.")

    add_para(doc, "Ключевые индикаторы платформы:", bold=True, space_after=3)
    add_bullet(doc,
        "Индикатор MCFTR/M2 — оценка рынка акций относительно денежной ликвидности.")
    add_bullet(doc,
        "Открытый интерес фьючерсов — позиции физических и юридических лиц "
        "в разбивке по инструментам (5м / 60м / дневные).")
    add_bullet(doc,
        "Market Breadth — доля акций, торгующихся выше скользящей средней: "
        "индикатор «здоровья» рынка.")
    add_bullet(doc,
        "Карта рынка (Heatmap) — визуализация изменений цен и оборотов "
        "акций по секторам.")
    add_bullet(doc,
        "СЧА фондов БПИФ — динамика стоимости чистых активов по категориям "
        "фондов (акции, облигации, золото, денежный рынок).")

    add_spacer(doc, 4)
    add_para(doc,
        "Подробное описание каждого индикатора с указанием используемых данных, "
        "источников и формул расчёта приведено в Приложении 1 к настоящему обращению.")

    add_para(doc, "Статус проекта:", bold=True, space_after=3)
    add_bullet(doc, "Стадия: активная разработка, предзапуск.")
    add_bullet(doc, "Команда: 2 человека (разработка и аналитика).")
    add_bullet(doc, "Юридическая форма: Индивидуальный предприниматель.")
    add_bullet(doc,
        "Продукт ориентирован исключительно на конечных пользователей — "
        "физических лиц. Перепродажа данных третьим лицам не предусмотрена "
        "и будет исключена условиями договора.")

    # ══════════════════════════════════════════════════════════
    # 3. ФИНАНСОВОЕ ОБОСНОВАНИЕ
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "3. Финансовое обоснование")

    add_para(doc,
        "Проект реализуется на собственные средства основателей без "
        "привлечения внешнего финансирования. На стадии до запуска и в "
        "первые месяцы после запуска команда несёт следующие расходы:")

    # Таблица расходов
    add_spacer(doc, 4)
    tbl = doc.add_table(rows=6, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    headers = ["Статья расходов", "В месяц", "В год"]
    header_bg = "2C4770"
    for ci, h in enumerate(headers):
        cell = tbl.rows[0].cells[ci]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold  = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.size  = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, header_bg)

    rows_data = [
        ("Серверная инфраструктура (VPS, БД)", "от [X] руб.", "от [X] руб."),
        ("Домен, SSL, CDN", "— ",            "[X] руб."),
        ("Прочие инструменты разработки",    "от [X] руб.", "от [X] руб."),
        ("Данные MOEX BI (стандартный тариф)", "от [X] руб.", "от [X] руб."),
        ("ИТОГО (без данных MOEX / с данными)", "— ",        "[X] / [X] руб."),
    ]
    alt_colors = ["F5F7FA", "FFFFFF"]
    for ri, (c1, c2, c3) in enumerate(rows_data, start=1):
        row = tbl.rows[ri]
        for ci, txt in enumerate([c1, c2, c3]):
            cell = row.cells[ci]
            cell.text = txt
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if ci > 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(cell, alt_colors[ri % 2])
        # выделить итоговую строку
        if ri == len(rows_data):
            for ci in range(3):
                tbl.rows[ri].cells[ci].paragraphs[0].runs[0].font.bold = True

    add_spacer(doc, 6)
    add_para(doc,
        "Стоимость данных по стандартному тарифу MOEX BI составляет значительную "
        "долю совокупных расходов проекта на стадии роста. При нулевой абонентской "
        "базе на старте полный тариф делает проект экономически нежизнеспособным "
        "в первые 12–18 месяцев. Именно поэтому льготный период является "
        "ключевым условием для запуска продукта.")

    add_para(doc,
        "Стоимость архивных данных исчисляется миллионами рублей, что "
        "при единовременной оплате недоступно для команды стартапа. "
        "Рассрочка на 36 месяцев позволит обеспечить полноценную "
        "историческую глубину данных с первого дня работы платформы.")

    # ══════════════════════════════════════════════════════════
    # 4. БИЗНЕС-МОДЕЛЬ И ПРОГНОЗ
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "4. Бизнес-модель и прогноз развития")

    add_para(doc, "Монетизация:", bold=True, space_after=3)
    add_bullet(doc,
        "Модель: платная подписка для физических лиц.")
    add_bullet(doc,
        "Тарифные планы: базовый и расширенный (конкретные цены "
        "определяются по результатам рыночного тестирования).")
    add_bullet(doc,
        "Данные конечным пользователям предоставляются исключительно "
        "в агрегированном аналитическом виде; сырые биржевые данные "
        "не передаются и не перепродаются.")

    add_spacer(doc, 4)
    add_para(doc, "Прогноз пользователей (консервативный сценарий):", bold=True, space_after=3)

    # Таблица прогноза
    tbl2 = doc.add_table(rows=4, cols=4)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl2.style = "Table Grid"

    h2 = ["Период", "Пользователей\n(платных)", "Среднемесячная\nвыручка", "Статус"]
    for ci, h in enumerate(h2):
        cell = tbl2.rows[0].cells[ci]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, header_bg)

    forecast = [
        (f"{YEAR} — {YEAR+1} (год 1)",     "0 → [X]",    "[X] руб.",  "Льготный период"),
        (f"{YEAR+1} — {YEAR+2} (год 2)",   "[X] → [X]",  "[X] руб.",  "Льготный период"),
        (f"{YEAR+2}+ (год 3+)",             "[X]+",        "[X]+ руб.", "Стандартный тариф"),
    ]
    for ri, row_data in enumerate(forecast, start=1):
        for ci, txt in enumerate(row_data):
            cell = tbl2.rows[ri].cells[ci]
            cell.text = txt
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_bg(cell, alt_colors[ri % 2])

    add_spacer(doc, 6)
    add_para(doc,
        "По окончании льготного периода (2 года) проект планирует перейти "
        "на стандартные коммерческие условия. Наличие абонентской базы "
        "и выручки к этому моменту обеспечит финансовую устойчивость "
        "для полноценных расчётов за данные.")

    # ══════════════════════════════════════════════════════════
    # 5. ВЫГОДА ДЛЯ МОСБИРЖИ
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "5. Выгода для ПАО Московская Биржа")

    add_para(doc,
        "Платформа «Фрейм» является инструментом экосистемы МосБиржи: "
        "она не конкурирует с биржей, а формирует дополнительную "
        "аудиторию активных частных инвесторов, торгующих на российском рынке.")

    add_bullet(doc,
        "Расширение аудитории. Аналитические инструменты платформы "
        "снижают порог входа для начинающих инвесторов и стимулируют "
        "более осознанное участие в торгах на МосБирже.")
    add_bullet(doc,
        "Рост активности. Пользователи, работающие с индикаторами "
        "(ОИ, Heatmap, Market Breadth), принимают более обоснованные "
        "торговые решения — что ведёт к росту числа сделок.")
    add_bullet(doc,
        "Популяризация срочного рынка. Отдельный блок по открытому "
        "интересу фьючерсов помогает розничным инвесторам разобраться "
        "в инструментах срочного рынка и начать с ними работать.")
    add_bullet(doc,
        "Обратная связь. Стартап в роли активного потребителя данных "
        "формирует полезную обратную связь по качеству и составу "
        "предоставляемых данных.")

    # ══════════════════════════════════════════════════════════
    # 6. ЗАПРАШИВАЕМЫЕ УСЛОВИЯ
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "6. Запрашиваемые условия")

    add_para(doc, "6.1. Льготный доступ к текущим данным", bold=True, space_after=3)
    add_bullet(doc,
        "Объём: данные MOEX BI по срочному рынку (FORTS), "
        "рынку акций (TQBR) и индексам.")
    add_bullet(doc,
        "Срок: 24 (двадцать четыре) месяца с даты подписания договора.")
    add_bullet(doc,
        "Стоимость в льготный период: 0 рублей.")
    add_bullet(doc,
        "По истечении льготного периода — переход на действующий "
        "коммерческий тариф МосБиржи.")

    add_spacer(doc, 4)
    add_para(doc, "6.2. Рассрочка на архивные данные", bold=True, space_after=3)
    add_bullet(doc,
        "Объём: исторические данные по срочному рынку, рынку акций и индексам "
        "(глубина и точный состав — по результатам переговоров).")
    add_bullet(doc,
        "Срок рассрочки: 36 (тридцать шесть) месяцев.")
    add_bullet(doc,
        "Периодичность платежей: ежемесячно, равными долями.")
    add_bullet(doc,
        "Первый платёж: через 24 месяца с даты договора "
        "(по окончании льготного периода по текущим данным).")

    add_spacer(doc, 6)
    add_para(doc,
        "Готовы рассмотреть альтернативные форматы сотрудничества, "
        "предложенные МосБиржей, в рамках программ поддержки финтех-стартапов "
        "и развития экосистемы биржи.")

    # ══════════════════════════════════════════════════════════
    # 7. ОБЯЗАТЕЛЬСТВА ПРОЕКТА
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "7. Обязательства со стороны проекта «Фрейм»")

    add_bullet(doc,
        "Не осуществлять перепродажу, передачу или трансляцию сырых "
        "биржевых данных третьим лицам в любой форме.")
    add_bullet(doc,
        "Использовать данные исключительно для расчёта аналитических "
        "индикаторов и отображения агрегированных результатов "
        "конечным пользователям-физлицам.")
    add_bullet(doc,
        "Указывать ПАО Московская Биржа как источник данных "
        "на платформе и в маркетинговых материалах.")
    add_bullet(doc,
        "Предоставлять МосБирже ежеквартальную отчётность "
        "о количестве пользователей и динамике роста в течение льготного периода.")
    add_bullet(doc,
        "Перейти на стандартный коммерческий тариф по истечении "
        "24-месячного льготного периода.")

    # ══════════════════════════════════════════════════════════
    # ЗАКЛЮЧЕНИЕ
    # ══════════════════════════════════════════════════════════
    add_heading(doc, "Заключение")

    add_para(doc,
        "Просим рассмотреть настоящее обращение и сообщить о возможности "
        "предоставления указанных льготных условий либо предложить "
        "альтернативные форматы сотрудничества.")

    add_para(doc,
        "Готовы предоставить дополнительные материалы, провести "
        "презентацию продукта или встретиться для обсуждения деталей "
        "в удобное для вас время.")

    add_spacer(doc, 14)

    sign = doc.add_paragraph()
    sign.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sign.paragraph_format.space_after = Pt(4)
    r = sign.add_run(
        f"С уважением,\n\n"
        f"Индивидуальный предприниматель\n"
        f"[ФИО]  _________________ /подпись/\n\n"
        f"«___» ____________ {YEAR} г."
    )
    r.font.size = Pt(10.5)

    add_spacer(doc, 14)

    # ══════════════════════════════════════════════════════════
    # ПРИЛОЖЕНИЕ
    # ══════════════════════════════════════════════════════════
    doc.add_page_break()

    app_title = doc.add_paragraph()
    app_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    app_title.paragraph_format.space_after = Pt(4)
    r = app_title.add_run("ПРИЛОЖЕНИЕ 1")
    r.font.size = Pt(12)
    r.font.bold = True

    app_sub = doc.add_paragraph()
    app_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    app_sub.paragraph_format.space_after = Pt(12)
    r = app_sub.add_run(
        "к обращению ИП [ФИО]\n"
        "об условиях доступа к данным для платформы «Фрейм»"
    )
    r.font.size = Pt(11)
    r.italic = True

    add_para(doc,
        "Описание аналитических индикаторов платформы «Фрейм» с указанием "
        "используемых данных, источников и методологии расчёта — "
        "см. файл «MOEX_Indicators_For_Submission_v2.xlsx», прилагаемый к настоящему обращению.",
        space_after=8)

    # Краткая сводная таблица индикаторов
    tbl3 = doc.add_table(rows=6, cols=3)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl3.style = "Table Grid"

    for ci, h in enumerate(["Индикатор", "Данные MOEX", "Таймфрейм"]):
        cell = tbl3.rows[0].cells[ci]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, header_bg)

    app_rows = [
        ("MCFTR / M2",        "Индекс MCFTR (ISS MOEX)",              "Ежедневно"),
        ("Открытый интерес",  "Данные BI по срочному рынку (FORTS)",   "5м / 60м / дневные"),
        ("Market Breadth",    "Свечи акций (BI, TQBR)",                "Дневные"),
        ("Карта рынка",       "Свечи акций (BI, TQBR)",                "Дневные"),
        ("СЧА фондов БПИФ",  "Данные БПИФ и индексы (ISS MOEX)",      "Ежедневно"),
    ]
    for ri, row_data in enumerate(app_rows, start=1):
        for ci, txt in enumerate(row_data):
            cell = tbl3.rows[ri].cells[ci]
            cell.text = txt
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            set_cell_bg(cell, alt_colors[ri % 2])

# ─────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # Базовый шрифт
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    build(doc)

    out = r"C:\Users\Vadim\PycharmProjects\MOEX\Фрейм_Обращение_МосБиржа.docx"
    doc.save(out)


if __name__ == "__main__":
    main()
